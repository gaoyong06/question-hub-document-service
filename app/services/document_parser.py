"""
文档解析服务
使用 python-docx 将 Word（.docx，或 .doc 先转为 .docx）转为 Markdown，供统一流水线解析。
其他格式通过 MarkItDown 转换为 Markdown 后解析。
"""
from __future__ import annotations

import os
import re
import subprocess
import tempfile
import zipfile
import uuid
from dataclasses import dataclass, field
from typing import List, Optional, Tuple
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from loguru import logger

from app.models import QuestionResult
from app.config import settings

# 格式辅助：正文默认字号（pt），用于判断「明显大于正文」的标题
DEFAULT_BODY_FONT_PT = 12.0
# 若段落字号 >= 此值且大于正文，可视为标题
SECTION_HEADING_FONT_PT_MIN = 14.0


def _format_suggests_section_heading(
    fmt: Optional[ParagraphFormatInfo],
    body_font_pt: float = DEFAULT_BODY_FONT_PT,
) -> bool:
    """根据段落格式判断是否像大题标题（辅助，与文字规则结合使用）。"""
    if fmt is None:
        return False
    style = (fmt.style_name or "").strip().lower()
    if "heading" in style or "标题" in style or style.startswith("heading") or style.startswith("标题"):
        return True
    if fmt.font_size_pt is not None and fmt.font_size_pt >= SECTION_HEADING_FONT_PT_MIN and fmt.font_size_pt > body_font_pt:
        return True
    if fmt.is_bold and fmt.left_indent_pt <= 0 and fmt.first_line_indent_pt <= 0:
        return True
    return False


# 试卷中常见大题标题 pattern：一. 二. 三. … 十. 或 一、二、
SECTION_HEADING_PATTERN = re.compile(
    r"^(一|二|三|四|五|六|七|八|九|十)[\.．、]\s*[^\s]"
)

# 题号/选项序号规范化：全角/半角点、顿号 → 半角 \. + 空格（转义点号，避免被 Markdown 解析为有序列表导致每题都显示为 1.）
# 必须包含半角 .，Word 中题号多为半角点（如 1.xxx），否则无法规范化
QUESTION_NUMBER_NORMALIZE = re.compile(r"^(\d+)[\.．、]\s*", re.MULTILINE)
OPTION_MARKER_NORMALIZE = re.compile(r"([A-Da-d])[\.．、]\s*")

# 缩进转无序列表：约每 18pt 一级，对应 Markdown 每级 2 空格
INDENT_PT_PER_LEVEL = 18.0

# 文件签名：DOCX 为 ZIP (PK)；旧版 Word .doc 为 OLE 复合文档 (D0 CF)
ZIP_DOCX_SIGNATURE = b"PK"
OLE_DOC_SIGNATURE = b"\xd0\xcf"


def _is_valid_word_signature(first_bytes: bytes) -> bool:
    """是否为支持的 Word 格式：.docx (ZIP) 或 .doc (OLE)。"""
    return first_bytes == ZIP_DOCX_SIGNATURE or first_bytes == OLE_DOC_SIGNATURE


def _ensure_doc_path_for_ole(local_path: str) -> str:
    """
    若文件内容为 .doc (OLE)，且当前路径为 .docx，则重命名为 .doc，便于后续 convert_doc_to_docx 识别。
    返回最终使用的路径。
    """
    if not os.path.isfile(local_path):
        return local_path
    with open(local_path, "rb") as f:
        sig = f.read(2)
    if sig != OLE_DOC_SIGNATURE:
        return local_path
    path_lower = local_path.lower()
    if not path_lower.endswith(".docx"):
        return local_path
    doc_path = os.path.splitext(local_path)[0] + ".doc"
    os.rename(local_path, doc_path)
    logger.info("Renamed OLE file to .doc for conversion: %s -> %s", local_path, doc_path)
    return doc_path


def _normalize_markdown_list_markers(text: str) -> str:
    """将题号/选项序号统一为「数字\\. / A\\. 」形式：转义点号避免 Markdown 有序列表语法，渲染时显示为 1. 2. 3. 而非全是 1.；exam_structure_utils 已支持 \\. 可选匹配。"""
    if not text:
        return text
    s = QUESTION_NUMBER_NORMALIZE.sub(r"\1\\. ", text)
    s = OPTION_MARKER_NORMALIZE.sub(r"\1\\. ", s)
    return s


def _indent_to_markdown_list_prefix(left_indent_pt: float) -> str:
    """根据 Word 左缩进返回 Markdown 无序列表前缀（  - /    - / ...）。"""
    if left_indent_pt <= 0:
        return ""
    level = max(1, int(left_indent_pt / INDENT_PT_PER_LEVEL) + 1)
    return "  " * level + "- "


@dataclass
class ParagraphFormatInfo:
    """Word 段落格式信息，用于辅助判断大题/小题（结构优先）。无格式时为 None。"""
    style_name: str = ""
    font_size_pt: Optional[float] = None
    left_indent_pt: float = 0.0
    first_line_indent_pt: float = 0.0
    is_bold: bool = False


class DocumentParser:
    """文档解析器（支持Word格式，其他格式通过MarkItDown转换）"""
    
    def __init__(self):
        self.temp_dir = Path(settings.temp_file_dir)
        self.temp_dir.mkdir(parents=True, exist_ok=True)
    
    def download_file(self, file_url: str) -> str:
        """
        下载文件到临时目录
        
        支持多种协议：
        - http://, https://: 通过HTTP下载
        - file://: 直接访问本地文件系统
        - 相对路径或绝对路径: 作为本地文件路径处理
        
        Args:
            file_url: 文件URL或路径
            
        Returns:
            本地文件路径
        """
        import httpx
        from urllib.parse import urlparse
        
        # 解析URL
        parsed = urlparse(file_url)
        scheme = parsed.scheme.lower()
        
        # 从URL提取文件名
        file_name = os.path.basename(parsed.path or file_url.split("?")[0])
        if not file_name or file_name == "content":
            # 如果文件名为空或是 "content"，使用 UUID 生成唯一文件名
            import uuid
            file_name = f"{uuid.uuid4().hex}.docx"
        else:
            # 确保文件扩展名
            if not any(file_name.lower().endswith(ext) for ext in settings.supported_extensions):
                file_name += ".docx"
        
        local_path = self.temp_dir / file_name
        
        logger.info(f"Downloading file from {file_url} to {local_path} (scheme: {scheme})")
        
        # 根据协议类型处理
        if scheme in ('http', 'https'):
            # 若 file_url 为对外网关 URL，走网关可能 401，后端应使用内网 asset-service 直连。
            # asset-service 主资源为 GET /asset/v1/files/{fileId}。
            download_url = file_url
            try:
                internal_parsed = urlparse(settings.asset_service_api_base_url)
                internal_host = (internal_parsed.netloc or "").strip()
                if internal_host and parsed.netloc != internal_host:
                    path = (parsed.path or "").strip("/")
                    if "files" in path and "asset/v1/files" in path:
                        parts = path.split("/")
                        if "files" in parts:
                            idx = parts.index("files")
                            if idx + 1 < len(parts):
                                file_id = parts[idx + 1]
                                if len(file_id) == 36 and file_id.count("-") == 4:
                                    base = settings.asset_service_api_base_url.rstrip("/")
                                    download_url = f"{base}/asset/v1/files/{file_id}"
                                    logger.info(f"Using internal asset URL for download: {download_url}")
            except Exception as e:
                logger.warning(f"Could not resolve internal asset URL, using original: {e}")
            # HTTP/HTTPS: 通过HTTP下载
            with httpx.Client(timeout=settings.download_timeout) as client:
                response = client.get(download_url)
                response.raise_for_status()
                
                # 检查响应类型
                content_type = response.headers.get("content-type", "").lower()
                logger.info(f"HTTP response Content-Type: {content_type}")
                
                # 如果是 JSON 响应（asset-service 的 DownloadFile 返回 JSON）
                if "application/json" in content_type:
                    import json
                    import base64
                    logger.info(f"Detected JSON response, parsing...")
                    data = response.json()
                    
                    # 提取文件数据（可能是 base64 编码的）
                    if isinstance(data, dict):
                        # 统一响应格式：{"success": true, "data": {...}}
                        if "data" in data:
                            file_data = data["data"]
                        else:
                            file_data = data
                        
                        # 如果 data 字段是字典，提取其中的 data 字段（base64 编码）
                        if isinstance(file_data, dict) and "data" in file_data:
                            # base64 解码
                            base64_str = file_data["data"]
                            logger.info(f"Decoding base64 data, length: {len(base64_str)}")
                            file_bytes = base64.b64decode(base64_str)
                            logger.info(f"Decoded file size: {len(file_bytes)} bytes")
                        elif isinstance(file_data, str):
                            # 直接是 base64 字符串
                            logger.info(f"Decoding base64 string, length: {len(file_data)}")
                            file_bytes = base64.b64decode(file_data)
                            logger.info(f"Decoded file size: {len(file_bytes)} bytes")
                        else:
                            raise ValueError(f"Unexpected JSON response format: {data}")
                    else:
                        raise ValueError(f"Unexpected JSON response format: {data}")
                    
                    # 检查文件大小
                    if len(file_bytes) > settings.max_file_size:
                        raise ValueError(f"File size {len(file_bytes)} exceeds maximum {settings.max_file_size}")
                    
                    # 验证文件签名：支持 .docx (ZIP/PK) 与 .doc (OLE/D0 CF)
                    if len(file_bytes) < 2 or not _is_valid_word_signature(file_bytes[:2]):
                        raise ValueError(
                            f"File is not a supported Word format (expected DOCX or DOC; signature: {file_bytes[:2]!r})"
                        )
                    logger.info(
                        f"File has valid Word signature: {'ZIP/DOCX' if file_bytes[:2] == ZIP_DOCX_SIGNATURE else 'OLE/DOC'}"
                    )
                    
                    # 保存文件
                    with open(local_path, "wb") as f:
                        f.write(file_bytes)
                    
                    # 验证文件已保存
                    if not os.path.exists(local_path):
                        raise FileNotFoundError(f"File was not saved: {local_path}")
                    
                    saved_size = os.path.getsize(local_path)
                    logger.info(f"File saved to {local_path}, size: {saved_size} bytes")
                    
                    # OLE (.doc) 时保存为 .doc 扩展名，便于后续 convert_doc_to_docx
                    local_path = Path(_ensure_doc_path_for_ole(str(local_path)))
                else:
                    # 直接是文件流
                    # 检查文件大小
                    content_length = response.headers.get("content-length")
                    if content_length and int(content_length) > settings.max_file_size:
                        raise ValueError(f"File size {content_length} exceeds maximum {settings.max_file_size}")
                    
                    # 保存文件
                    with open(local_path, "wb") as f:
                        for chunk in response.iter_bytes(chunk_size=settings.download_chunk_size):
                            f.write(chunk)
                    
                    # 验证文件已保存
                    if not os.path.exists(local_path):
                        raise FileNotFoundError(f"File was not saved: {local_path}")
                    
                    saved_size = os.path.getsize(local_path)
                    logger.info(f"File saved to {local_path}, size: {saved_size} bytes")
                    
                    # 验证文件签名并统一 .doc 路径（OLE 时改为 .doc 供 convert_doc_to_docx）
                    with open(local_path, "rb") as f:
                        first_bytes = f.read(2)
                    if not _is_valid_word_signature(first_bytes):
                        raise ValueError(
                            f"Saved file is not a supported Word format (signature: {first_bytes!r})"
                        )
                    local_path = Path(_ensure_doc_path_for_ole(str(local_path)))
        
        elif scheme == 'file':
            # file:// 协议: 直接访问本地文件
            # file:///path/to/file 或 file://localhost/path/to/file
            file_path = parsed.path
            # 处理 Windows 路径 (file:///C:/path/to/file)
            if os.name == 'nt' and len(file_path) > 2 and file_path[0] == '/' and file_path[2] == ':':
                file_path = file_path[1:]  # 移除开头的 /
            
            if not os.path.isabs(file_path):
                raise ValueError(f"file:// URL must be an absolute path: {file_path}")
            
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # 检查文件大小
            file_size = os.path.getsize(file_path)
            if file_size > settings.max_file_size:
                raise ValueError(f"File size {file_size} exceeds maximum {settings.max_file_size}")
            
            # 复制文件到临时目录
            import shutil
            shutil.copy2(file_path, local_path)
            logger.info(f"Copied file from {file_path} to {local_path}")
            # 若为 .doc (OLE)，统一为 .doc 扩展名
            local_path = Path(_ensure_doc_path_for_ole(str(local_path)))
        
        else:
            # 无协议或未知协议: 作为本地文件路径处理
            # 可能是相对路径或绝对路径
            file_path = file_url
            
            # 如果是相对路径，尝试从常见位置查找
            if not os.path.isabs(file_path):
                # 尝试从 asset-service 的默认存储路径查找
                possible_paths = [
                    file_path,  # 原始路径
                    os.path.join("./uploads", file_path),  # 相对 uploads 目录
                    os.path.join("/uploads", file_path),  # 绝对 uploads 目录
                ]
                
                found = False
                for path in possible_paths:
                    abs_path = os.path.abspath(path)
                    if os.path.exists(abs_path):
                        file_path = abs_path
                        found = True
                        break
                
                if not found:
                    raise FileNotFoundError(f"File not found: {file_url} (tried: {possible_paths})")
            else:
                if not os.path.exists(file_path):
                    raise FileNotFoundError(f"File not found: {file_path}")
            
            # 检查文件大小
            file_size = os.path.getsize(file_path)
            if file_size > settings.max_file_size:
                raise ValueError(f"File size {file_size} exceeds maximum {settings.max_file_size}")
            
            # 复制文件到临时目录
            import shutil
            shutil.copy2(file_path, local_path)
            logger.info(f"Copied file from {file_path} to {local_path}")
            local_path = Path(_ensure_doc_path_for_ole(str(local_path)))
        
        # 最终验证：确保文件存在且有效
        if not os.path.exists(local_path):
            raise FileNotFoundError(f"File was not saved: {local_path}")
        
        file_size = os.path.getsize(local_path)
        if file_size == 0:
            raise ValueError(f"Downloaded file is empty: {local_path}")
        
        # 验证文件签名（支持 DOCX 与 DOC），并统一 .doc 路径
        with open(local_path, "rb") as f:
            first_bytes = f.read(2)
        if not _is_valid_word_signature(first_bytes):
            raise ValueError(
                f"Downloaded file is not a supported Word format (signature: {first_bytes!r}, size: {file_size})"
            )
        local_path = _ensure_doc_path_for_ole(str(local_path))
        logger.info(
            f"File downloaded successfully: {local_path}, size: {os.path.getsize(local_path)} bytes"
        )
        return local_path

    def convert_doc_to_docx(self, doc_path: str) -> Optional[str]:
        """
        将 .doc 转为 .docx（依赖系统 LibreOffice soffice），便于后续用 python-docx 转 Markdown。

        Args:
            doc_path: 本地 .doc 文件路径

        Returns:
            转换后的 .docx 路径；若未安装 soffice 或转换失败则返回 None，调用方可用 MarkItDown 处理 .doc。
        """
        if not os.path.isfile(doc_path) or not doc_path.lower().endswith(".doc"):
            return None
        out_dir = os.path.dirname(doc_path)
        try:
            subprocess.run(
                ["soffice", "--headless", "--convert-to", "docx", "--outdir", out_dir, doc_path],
                check=True,
                capture_output=True,
                timeout=60,
            )
            docx_path = os.path.splitext(doc_path)[0] + ".docx"
            if os.path.isfile(docx_path):
                logger.info("Converted .doc to .docx: %s", docx_path)
                return docx_path
        except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired) as e:
            logger.warning("soffice convert .doc to .docx failed (use MarkItDown for .doc): %s", e)
        return None

    def docx_to_markdown(self, file_path: str) -> Tuple[str, dict]:
        """
        从 .docx 文件生成 Markdown 字符串（段落间 \\n\\n，段内保留 {{IMAGE_N}} 替换为 ![](本地绝对路径)），
        供统一走 process_images_in_markdown 上传图片后得到最终 markdown_content。

        Args:
            file_path: 本地 .docx 文件路径（必须是 OOXML，即 .docx）

        Returns:
            (markdown_content, metadata)，metadata 含 title 等，与 MarkItDown 返回格式对齐。
        """
        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"File does not exist: {file_path}")
        with open(file_path, "rb") as f:
            if f.read(2) != b"PK":
                raise ValueError("File is not a valid ZIP/DOCX file")
        doc = Document(file_path)
        document_title = (getattr(doc.core_properties, "title", None) or "").strip()
        paragraphs, image_paths, format_infos = self._extract_paragraphs_with_images(doc, file_path)
        if not document_title and paragraphs:
            for p in paragraphs:
                t = (p or "").strip()
                if t and not t.startswith("{{IMAGE_"):
                    document_title = t[:200].strip()
                    break
        # 根据段落格式转为 Markdown：标题层级（# / ##）、缩进→无序列表（  - ）、题号/选项序号规范化
        body_font_pt = DEFAULT_BODY_FONT_PT
        lines = []
        first_content_idx = None
        for i, p in enumerate(paragraphs):
            t = (p or "").strip()
            if not t:
                lines.append("")
                continue
            if first_content_idx is None and not t.startswith("{{IMAGE_"):
                first_content_idx = i
            fmt = format_infos[i] if i < len(format_infos) else None
            heading_prefix = self._markdown_heading_prefix_for_paragraph(
                t, fmt, body_font_pt, is_first_content=(i == first_content_idx)
            )
            s = p
            # 先做题号/选项序号规范化（仅对占位符段落），再替换为图片路径；否则路径中的 "x.png" 会被 OPTION_MARKER_NORMALIZE 误匹配（如 c9cb 的 b.）变成 "x. png" 导致上传时文件不存在
            s = _normalize_markdown_list_markers(s)
            for j, path in enumerate(image_paths):
                s = s.replace("{{IMAGE_%d}}" % j, "![](%s)" % (path,))
            # 标题前缀 或 缩进→无序列表前缀（无标题且 left_indent_pt > 0）
            if heading_prefix:
                s = heading_prefix + s.strip()
            else:
                list_prefix = _indent_to_markdown_list_prefix(fmt.left_indent_pt) if fmt else ""
                if list_prefix:
                    s = list_prefix + s.strip()
            lines.append(s)
        markdown_content = "\n\n".join(lines)
        metadata = {"title": document_title}
        logger.info("docx_to_markdown: %s chars, %s images", len(markdown_content), len(image_paths))
        return markdown_content, metadata

    def _markdown_heading_prefix_for_paragraph(
        self,
        paragraph_text: str,
        fmt: Optional[ParagraphFormatInfo],
        body_font_pt: float,
        is_first_content: bool,
    ) -> str:
        """
        根据段落文本、样式与顺序返回 Markdown 标题前缀。
        优先依据文档样式（字号、加粗、Word 样式）和顺序（首段），其次才用文字 pattern。
        - 文档标题：首段 + 标题样貌（Title 样式 / 字号明显大于正文 / 加粗）+ 非「一. 二.」→ "# "
        - 大题/节标题：Heading 样式、或「一. 二.」、或格式像大题且非首段 → "## "
        - 其余 → ""
        """
        style = (fmt.style_name or "").strip().lower() if fmt else ""
        is_section_text = bool(paragraph_text and SECTION_HEADING_PATTERN.match(paragraph_text.strip()))
        format_like_heading = _format_suggests_section_heading(fmt, body_font_pt)

        # 1. Word 明确标题样式（与是否首段无关）
        if "title" in style or style == "标题":
            return "# "

        # 2. 文字明确为大题（一. 二. …）→ 二级标题
        if is_section_text and (not fmt or fmt.left_indent_pt <= 0):
            return "## "

        # 3. 格式像标题（字号大、加粗、Heading 样式等）：结合顺序区分「文档标题」与「大题」
        if format_like_heading:
            # 首段且不是「一. 二.」→ 视为试卷总标题（通常字号更大、在首行）
            if is_first_content and paragraph_text and not paragraph_text.startswith("{{IMAGE_") and not is_section_text:
                return "# "
            return "## "

        # 4. 无明确样式时：仅当首段 + 较短 + 有标题特征（加粗或字号大于正文）+ 非大题文字 → 文档标题
        if is_first_content and paragraph_text and not paragraph_text.startswith("{{IMAGE_") and not is_section_text:
            if len(paragraph_text) < 100 and (not fmt or (fmt.is_bold or (fmt.font_size_pt and fmt.font_size_pt > body_font_pt))):
                return "# "

        # 5. 兜底：仅文字像大题（无格式时）
        if is_section_text:
            return "## "
        return ""

    def _get_paragraph_format(self, para) -> ParagraphFormatInfo:
        """从 python-docx 段落对象提取格式信息，用于格式辅助结构判断。"""
        fmt = ParagraphFormatInfo()
        try:
            fmt.style_name = getattr(para.style, "name", None) or ""
            pf = getattr(para, "paragraph_format", None)
            if pf is not None:
                li = getattr(pf, "left_indent", None)
                fmt.left_indent_pt = float(getattr(li, "pt", 0) or 0)
                fl = getattr(pf, "first_line_indent", None)
                fmt.first_line_indent_pt = float(getattr(fl, "pt", 0) or 0)
            font_sizes: List[float] = []
            bold_count = 0
            text_run_count = 0
            for run in getattr(para, "runs", []):
                if getattr(run, "text", "").strip():
                    text_run_count += 1
                    if getattr(run, "bold", None) is True:
                        bold_count += 1
                    sz = getattr(run.font, "size", None)
                    if sz is not None and hasattr(sz, "pt"):
                        font_sizes.append(float(sz.pt))
            if font_sizes:
                fmt.font_size_pt = max(font_sizes)
            fmt.is_bold = text_run_count > 0 and bold_count >= (text_run_count + 1) // 2
        except Exception as e:
            logger.debug("Failed to get paragraph format: %s", e)
        return fmt

    def _extract_paragraphs_with_images(self, doc: Document, file_path: str) -> Tuple[List[str], List[str], List[ParagraphFormatInfo]]:
        """
        提取段落文本，并将文档中的内嵌图片导出为临时文件，在文本中用 {{IMAGE_N}} 占位。
        同时提取每段格式信息（样式、字号、缩进、加粗），供结构解析辅助判断大题/小题。
        返回 (段落列表, 图片路径列表, 每段对应的格式信息列表，与段落一一对应)。
        """
        image_paths: List[str] = []
        image_dir = self.temp_dir / "docx_images"
        image_dir.mkdir(parents=True, exist_ok=True)

        def save_image_blob(blob: bytes, ext: str = ".png") -> str:
            path = image_dir / f"{uuid.uuid4().hex}{ext}"
            path.write_bytes(blob)
            return str(path)

        w_br = qn("w:br")
        w_r = qn("w:r")
        w_t = qn("w:t")

        paragraphs: List[str] = []
        format_infos: List[ParagraphFormatInfo] = []
        for para in doc.paragraphs:
            fmt = self._get_paragraph_format(para)
            parts: List[str] = []
            run_idx = 0
            for child in para._element:
                if child.tag == w_br:
                    parts.append("\n")
                    continue
                if child.tag != w_r:
                    continue
                run = para.runs[run_idx] if run_idx < len(para.runs) else None
                run_idx += 1
                if run is None:
                    for t in child.iter():
                        if t.tag == w_t and t.text:
                            parts.append(t.text)
                        if t.tail:
                            parts.append(t.tail)
                    continue
                if run._element is None:
                    if run.text:
                        parts.append(run.text)
                    continue
                r_id = None
                drawing = run._element.find(qn("w:drawing"))
                if drawing is not None:
                    blip = None
                    for elem in drawing.iter():
                        if elem.tag is not None and "blip" in elem.tag.lower():
                            blip = elem
                            break
                    if blip is None:
                        blip = drawing.find(qn("a:blip"))
                    if blip is not None:
                        r_id = blip.get(qn("r:embed"))
                if r_id is None:
                    pict = run._element.find(qn("w:pict"))
                    if pict is not None:
                        for elem in pict.iter():
                            if elem.tag is not None and "imagedata" in elem.tag.lower():
                                r_id = elem.get(qn("r:id"))
                                if r_id:
                                    break
                if r_id and hasattr(doc.part, "related_parts") and r_id in doc.part.related_parts:
                    part = doc.part.related_parts[r_id]
                    blob = getattr(part, "blob", None) or getattr(part, "_blob", None)
                    if blob is not None:
                        ext = ".png"
                        if getattr(part, "content_type", "").startswith("image/jpeg") or getattr(
                            part, "partname", ""
                        ).lower().endswith(".emf"):
                            ext = ".jpg"
                        local_path = save_image_blob(blob, ext)
                        image_paths.append(local_path)
                        parts.append("{{IMAGE_%d}}" % (len(image_paths) - 1))
                        continue
                if run.text:
                    parts.append(run.text)
            line = "".join(parts).strip()
            if line:
                paragraphs.append(line)
                format_infos.append(fmt)

        if not image_paths and doc.paragraphs:
            try:
                with zipfile.ZipFile(file_path, "r") as zf:
                    names = [n for n in zf.namelist() if n.startswith("word/media/")]
                    for i, name in enumerate(sorted(names)):
                        data = zf.read(name)
                        ext = os.path.splitext(name)[1] or ".png"
                        local_path = str(image_dir / f"img_{i}{ext}")
                        Path(local_path).write_bytes(data)
                        image_paths.append(local_path)
            except Exception as e:
                logger.debug("Fallback zip media extraction failed: %s", e)

        return paragraphs, image_paths, format_infos

    def cleanup(self, file_path: str):
        """清理临时文件"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"Cleaned up temporary file: {file_path}")
        except Exception as e:
            logger.warning(f"Failed to cleanup file {file_path}: {e}")

